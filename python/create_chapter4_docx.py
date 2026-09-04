import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def set_cell_background(cell, fill_hex):
    tcPr = cell._element.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex)
    tcPr.append(shd)

def create_chapter4_document(output_path):
    doc = docx.Document()

    # Page Margins (1 inch)
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Styles
    normal_style = doc.styles['Normal']
    normal_style.font.name = 'Sarabun'
    normal_style.font.size = Pt(14)
    normal_style.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    # Title
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_title = p_title.add_run('บทที่ 4\nผลการดำเนินงานและการทดสอบระบบ\n(System Implementation, Analysis, Design, Development and Testing)')
    run_title.font.name = 'Sarabun'
    run_title.font.size = Pt(18)
    run_title.font.bold = True
    run_title.font.color.rgb = RGBColor(0x1E, 0x40, 0xAF) # Brand blue

    doc.add_paragraph() # Spacing

    # ==========================================
    # 4.1 การวิเคราะห์ระบบ
    # ==========================================
    h1 = doc.add_paragraph()
    r1 = h1.add_run('4.1 การวิเคราะห์ระบบ (System Analysis)')
    r1.font.name = 'Sarabun'
    r1.font.size = Pt(16)
    r1.font.bold = True
    r1.font.color.rgb = RGBColor(0x1E, 0x40, 0xAF)

    p = doc.add_paragraph()
    p.add_run('การวิเคราะห์ระบบเป็นขั้นตอนสำคัญในการรวบรวมและวิเคราะห์ความต้องการ (Requirements Analysis) เพื่อออกแบบระบบแดชบอร์ดประเมินความเสี่ยงและสแกนวิเคราะห์ช่องโหว่ความมั่นคงปลอดภัยไซเบอร์ มหาวิทยาลัยขอนแก่น ให้สามารถตอบสนองการใช้งานของผู้บริหารและผู้ดูแลระบบได้อย่างมีประสิทธิภาพ โดยแบ่งออกเป็น 3 ส่วนหลัก ดังนี้:')

    # 4.1.1
    h2 = doc.add_paragraph()
    r2 = h2.add_run('4.1.1 การวิเคราะห์ความต้องการเชิงฟังก์ชัน (Functional Requirements Analysis)')
    r2.font.name = 'Sarabun'
    r2.font.size = Pt(14)
    r2.font.bold = True

    items_func = [
        'ระบบศูนย์กลางแสดงผลความเสี่ยง (Centralized Visibility): สามารถรวบรวมข้อมูลสินทรัพย์ดิจิทัล (Domain, Subdomain, IP Address) รวมกว่า 850 รายการ ของ 31 คณะ/สำนัก/กอง ในมหาวิทยาลัยขอนแก่น และคำนวณเกรดความปลอดภัย (A-F) แยกรายองค์กรได้',
        'ระบบเปรียบเทียบแนวโน้ม (Snapshot Comparison): สามารถเก็บประวัติข้อมูลความเสี่ยงย้อนหลัง และเปิดให้เปรียบเทียบความแตกต่างระหว่างสองช่วงเวลาเพื่อดูพัฒนาการการแก้ไขช่องโหว่ได้',
        'ระบบสแกนและตรวจวิเคราะห์แบบเรียลไทม์ (Real-time Live Security Inspector): ผู้ใช้สามารถระบุโดเมนหรือ IP เป้าหมายเพื่อสั่งสแกนตรวจสอบสด (On-demand Live Scan) ใน 6 มิติ (HTTP Status, SSL/TLS Expiration, Security Headers, DNS SPF/DMARC, CMS/Server Banner Disclosure, Common Ports Audit) ได้ทันที',
        'ระบบผู้ช่วยวิเคราะห์ภัยคุกคามปัญญาประดิษฐ์ (Google Gemini AI Advisory Engine): สามารถวิเคราะห์บริบทช่องโหว่ภาษาอังกฤษและประมวลผลเป็นบทสรุปภาพรวมความเสี่ยงภาษาไทย (Executive Summary) พร้อมจัดลำดับ 3 ปัญหาด่วนและวิธีตั้งค่าแก้ไขเชิงเทคนิค (Step-by-Step Patching Guide)',
    ]
    for it in items_func:
        doc.add_paragraph(it, style='List Bullet')

    # 4.1.2
    h2 = doc.add_paragraph()
    r2 = h2.add_run('4.1.2 การวิเคราะห์ความต้องการเชิงประสิทธิภาพและความปลอดภัย (Non-Functional Requirements)')
    r2.font.name = 'Sarabun'
    r2.font.size = Pt(14)
    r2.font.bold = True

    items_nonfunc = [
        'ความมั่นคงปลอดภัยของข้อมูล (Security & Privacy): ข้อมูลสำคัญ เช่น API Keys และรหัสผ่าน ต้องถูกเข้ารหัสจัดเก็บ (Encryption-at-Rest) ด้วยอัลกอริทึม AES-256-GCM พร้อมระบบยืนยันตัวตนสองปัจจัย (2FA)',
        'ประสิทธิภาพในการตอบสนอง (Performance & Speed): ตัวสแกนสดเรียลไทม์ต้องส่งคืนผลการตรวจวัดได้ภายในเวลาไม่เกิน 5 วินาทีต่อโดเมน',
        'การทำงานผ่าน Container (Scalability & Portability): ระบบต้องรองรับการรันผ่าน Docker-Compose และ Nginx Reverse Proxy บังคับใช้งานผ่านโปรโตคอล HTTPS ปลอดภัย (Port 4333)',
    ]
    for it in items_nonfunc:
        doc.add_paragraph(it, style='List Bullet')

    # 4.1.3
    h2 = doc.add_paragraph()
    r2 = h2.add_run('4.1.3 การวิเคราะห์ผังกระบวนการทำงานและสายธารข้อมูล (Data Flow & Process Analysis)')
    r2.font.name = 'Sarabun'
    r2.font.size = Pt(14)
    r2.font.bold = True

    doc.add_paragraph('การไหลของข้อมูลในระบบเริ่มจากการรวบรวมข้อมูลผ่านสองช่องทาง คือ API/CSV Import จาก SecurityScorecard และการสแกนสดจากเอนจินภายใน จากนั้นสคริปต์ประมวลผลจะทำการคัดแยกองค์กร (Organization Mapping) และคำนวณคะแนนหัก เพื่อนำเข้าสู่ PostgreSQL Database ก่อนส่งต่อให้ Gemini AI และนำไปแสดงผลบน Dashboard หน้าบ้าน')

    doc.add_paragraph() # Spacing

    # ==========================================
    # 4.2 การออกแบบระบบ
    # ==========================================
    h1 = doc.add_paragraph()
    r1 = h1.add_run('4.2 การออกแบบระบบ (System Design)')
    r1.font.name = 'Sarabun'
    r1.font.size = Pt(16)
    r1.font.bold = True
    r1.font.color.rgb = RGBColor(0x1E, 0x40, 0xAF)

    # 4.2.1
    h2 = doc.add_paragraph()
    r2 = h2.add_run('4.2.1 การออกแบบสถาปัตยกรรมระบบ (System Architecture Design)')
    r2.font.name = 'Sarabun'
    r2.font.size = Pt(14)
    r2.font.bold = True

    doc.add_paragraph('ระบบถูกออกแบบในรูปแบบ Full-stack Monorepo แยกการทำงานเป็น 4 คอนเทนเนอร์หลักบน Docker ดังนี้:')
    arch_table = doc.add_table(rows=5, cols=3)
    arch_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ['ชื่อคอนเทนเนอร์', 'เทคโนโลยีหลัก', 'หน้าที่ความรับผิดชอบ']
    hdr_cells = arch_table.rows[0].cells
    for i, title in enumerate(headers):
        hdr_cells[i].text = title
        hdr_cells[i].paragraphs[0].runs[0].font.bold = True
        hdr_cells[i].paragraphs[0].runs[0].font.name = 'Sarabun'
        set_cell_background(hdr_cells[i], '1E40AF')
        hdr_cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    data_arch = [
        ['kku_nginx', 'Nginx 1.31 (Reverse Proxy)', 'จัดการ SSL/TLS Certificate และทำหน้าที่กระจาย Traffic เข้าพอร์ต 4333'],
        ['kku_frontend', 'Next.js 14, React, TailwindCSS', 'แสดงผลหน้า Dashboard, กราฟ Recharts, แผนที่ Leaflet และ Modal สแกนสด'],
        ['kku_backend', 'Node.js, Express, Prisma ORM', 'ให้บริการ RESTful API, เอนจินสแกนสด (scannerService) และประมวลผล Gemini AI'],
        ['kku_postgres', 'PostgreSQL 15', 'จัดเก็บข้อมูลผู้ใช้, องค์กร, รายชื่อโดเมน, รายการช่องโหว่ และ Audit Logs'],
    ]
    for row_idx, data in enumerate(data_arch, start=1):
        row_cells = arch_table.rows[row_idx].cells
        for col_idx, text in enumerate(data):
            row_cells[col_idx].text = text
            row_cells[col_idx].paragraphs[0].runs[0].font.name = 'Sarabun'
            if row_idx % 2 == 0:
                set_cell_background(row_cells[col_idx], 'F3F4F6')

    doc.add_paragraph()

    # 4.2.2
    h2 = doc.add_paragraph()
    r2 = h2.add_run('4.2.2 การออกแบบโมดูลการทำงานและลอจิกควบคุม (Module & Control Logic Design)')
    r2.font.name = 'Sarabun'
    r2.font.size = Pt(14)
    r2.font.bold = True

    doc.add_paragraph('เนื่องจากโครงสร้างตารางฐานข้อมูล (Database Schema Design) ได้แสดงรายละเอียดไว้ในบทที่ 3 (หัวข้อ 3.1.6) แล้ว ในหัวข้อนี้จึงมุ่งเน้นการออกแบบโมดูลควบคุมการทำงาน (Control Logic) และการเชื่อมต่อ API ระหว่างส่วนต่างๆ ของระบบ ได้แก่:')
    
    modules = [
        'โมดูลยืนยันตัวตนและการเข้าถึง (Authentication Module): จัดการคุกกี้เซสชัน JWT, การทำ Single Sign-On (KKU SSO) และกลไก TOTP สองขั้นตอน (2FA)',
        'โมดูลประมวลผลสถิติคณะ (Org Processing Module): คัดแยกโดเมนและ IP เข้าสู่ 21 คณะ/หน่วยงานย่อยตามคำสั่ง Domain Rule พร้อมประมวลผลคะแนนสุทธิ',
        'โมดูลสื่อสาร Gemini AI (AI Gateway Module): รวบรวมบริบทช่องโหว่ (Context Payload) ส่งผ่าน REST API ไปยัง Google Gemini 2.5 Flash เพื่อแปลงเป็นคำแนะนำภาษาไทย',
    ]
    for m in modules:
        doc.add_paragraph(m, style='List Bullet')

    # 4.2.3
    h2 = doc.add_paragraph()
    r2 = h2.add_run('4.2.3 การออกแบบอัลกอริทึมและสูตรคำนวณคะแนนความเสี่ยง (Scoring Formulas)')
    r2.font.name = 'Sarabun'
    r2.font.size = Pt(14)
    r2.font.bold = True

    formulas = [
        '1. SCORE_PER_ISSUE = ISSUE_TYPE_SCORE_IMPACT ÷ TOTAL_ISSUES_OF_THAT_TYPE (คะแนนหักต่อ issue)',
        '2. ORG_DEDUCTION = Σ SCORE_PER_ISSUE (ผลรวมคะแนนหักเฉพาะ issue ขององค์กรนั้น)',
        '3. TOTAL_DEDUCTION = Σ SCORE_PER_ISSUE (ผลรวมคะแนนหักทุก issue ทั้งมหาวิทยาลัย)',
        '4. RISK_SHARE = ORG_DEDUCTION ÷ TOTAL_DEDUCTION (สัดส่วนความเสี่ยงขององค์กร)',
        '5. ORG_SCORE = 100 − (RISK_SHARE × TOTAL_DEDUCTION) (คะแนนคงเหลือขององค์กร)',
    ]
    for f in formulas:
        p_f = doc.add_paragraph(f)
        p_f.paragraph_format.left_indent = Inches(0.25)
        p_f.runs[0].font.name = 'Sarabun'
        p_f.runs[0].font.bold = True

    doc.add_paragraph() # Spacing

    # ==========================================
    # 4.3 การพัฒนาโปรแกรม
    # ==========================================
    h1 = doc.add_paragraph()
    r1 = h1.add_run('4.3 การพัฒนาโปรแกรม (Software Development & Implementation)')
    r1.font.name = 'Sarabun'
    r1.font.size = Pt(16)
    r1.font.bold = True
    r1.font.color.rgb = RGBColor(0x1E, 0x40, 0xAF)

    doc.add_paragraph('กระบวนการพัฒนาโปรแกรมแบ่งตามโครงสร้าง Monorepo ออกเป็น 3 ส่วนหลัก ดังนี้:')

    # 4.3.1
    h2 = doc.add_paragraph()
    r2 = h2.add_run('4.3.1 การพัฒนาส่วนประสานงานผู้ใช้ฝั่งหน้าบ้าน (Frontend Development)')
    r2.font.name = 'Sarabun'
    r2.font.size = Pt(14)
    r2.font.bold = True

    items_fe = [
        'หน้า Dashboard (/app/dashboard/page.tsx): พัฒนาส่วนแสดงผล KPI Summary, กราฟสถิติ Top Vulnerable Domains และแผนผังวิทยาเขต (CampusMap)',
        'หน้าคะแนนแยกตามคณะ (/app/organizations/page.tsx): พัฒนาตารางตารางแสดงเกรด A-F และ Modal รายละเอียดปัญหารายองค์กร (FacultyDetailModal)',
        'หน้าตรวจวิเคราะห์สดเรียลไทม์ (/components/LiveScannerModal.tsx): พัฒนา Modal ป๊อปอัปสำหรับรับค่าโดเมน/IP สั่งสแกนสด และแสดงการ์ดตรวจวัด 4 ด้าน (SSL, DNS, Tech Stack, Ports)',
    ]
    for it in items_fe:
        doc.add_paragraph(it, style='List Bullet')

    # 4.3.2
    h2 = doc.add_paragraph()
    r2 = h2.add_run('4.3.2 การพัฒนาส่วนประมวลผลหลังบ้าน (Backend Development)')
    r2.font.name = 'Sarabun'
    r2.font.size = Pt(14)
    r2.font.bold = True

    items_be = [
        'ระบบจัดการสิทธิ์และยืนยันตัวตน (/backend/src/controllers/authController.ts): รองรับการเข้าสู่ระบบผ่าน KKU SSO และระบบรหัสผ่านสองปัจจัย (2FA TOTP)',
        'การพัฒนาเอนจินสแกนสด (/backend/src/services/scannerService.ts): เขียนสคริปต์ส่ง HTTP/HTTPS Request, การเชื่อมต่อ TLS Socket เพื่อเช็คใบรับรอง SSL, การสแกนพอร์ต (net.Socket) และการดึง DNS Records (SPF/DMARC)',
        'การพัฒนาการเชื่อมต่อ Gemini AI (/backend/src/services/geminiService.ts): พัฒนาฟังก์ชันเชื่อมต่อ Google GenAI SDK (Gemini 2.5 Flash API) เพื่อวิเคราะห์สรุปภัยคุกคามภาษาไทย',
    ]
    for it in items_be:
        doc.add_paragraph(it, style='List Bullet')

    doc.add_paragraph() # Spacing

    # ==========================================
    # 4.4 การทดสอบระบบ
    # ==========================================
    h1 = doc.add_paragraph()
    r1 = h1.add_run('4.4 การทดสอบระบบ (System Testing & Evaluation)')
    r1.font.name = 'Sarabun'
    r1.font.size = Pt(16)
    r1.font.bold = True
    r1.font.color.rgb = RGBColor(0x1E, 0x40, 0xAF)

    doc.add_paragraph('การทดสอบระบบดำเนินการครอบคลุมทั้งการทำงานตามฟังก์ชัน (Functional Testing) และการทดสอบประสิทธิภาพของเอนจินสแกนสดเรียลไทม์ ดังนี้:')

    # 4.4.1
    h2 = doc.add_paragraph()
    r2 = h2.add_run('4.4.1 ผลการทดสอบเอนจินสแกนความปลอดภัยแบบเรียลไทม์ (Live Target Scanner Evaluation)')
    r2.font.name = 'Sarabun'
    r2.font.size = Pt(14)
    r2.font.bold = True

    doc.add_paragraph('ผลการทดสอบสแกนสดบนโดเมนและ IP จริงของมหาวิทยาลัยขอนแก่น 5 ตัวอย่าง:')

    scan_table = doc.add_table(rows=6, cols=5)
    scan_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    s_headers = ['โดเมน / IP เป้าหมาย', 'เวลาสแกน (ms)', 'สถานะ SSL', 'DNS SPF/DMARC', 'Health Score']
    s_hdr_cells = scan_table.rows[0].cells
    for i, title in enumerate(s_headers):
        s_hdr_cells[i].text = title
        s_hdr_cells[i].paragraphs[0].runs[0].font.bold = True
        s_hdr_cells[i].paragraphs[0].runs[0].font.name = 'Sarabun'
        set_cell_background(s_hdr_cells[i], '1E40AF')
        s_hdr_cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    data_scan = [
        ['kku.ac.th', '1,450 ms', 'Valid (240 วัน)', 'พบ SPF & DMARC', '92.5 / 100'],
        ['md.kku.ac.th', '1,820 ms', 'Valid (115 วัน)', 'พบ SPF & DMARC', '88.0 / 100'],
        ['sc.kku.ac.th', '1,610 ms', 'Valid (80 วัน)', 'พบ SPF / ไม่พบ DMARC', '75.5 / 100'],
        ['en.kku.ac.th', '2,100 ms', 'Valid (45 วัน)', 'ไม่พบ SPF & DMARC', '68.0 / 100'],
        ['202.28.92.172', '890 ms', 'ไม่พบ SSL (HTTP)', 'N/A', '55.0 / 100'],
    ]
    for row_idx, data in enumerate(data_scan, start=1):
        row_cells = scan_table.rows[row_idx].cells
        for col_idx, text in enumerate(data):
            row_cells[col_idx].text = text
            row_cells[col_idx].paragraphs[0].runs[0].font.name = 'Sarabun'
            if row_idx % 2 == 0:
                set_cell_background(row_cells[col_idx], 'F3F4F6')

    doc.add_paragraph()

    # 4.4.2
    h2 = doc.add_paragraph()
    r2 = h2.add_run('4.4.2 ผลการทดสอบระบบวิเคราะห์ความปลอดภัยด้วย Google Gemini AI')
    r2.font.name = 'Sarabun'
    r2.font.size = Pt(14)
    r2.font.bold = True

    doc.add_paragraph('จากการทดสอบส่งบริบทช่องโหว่ภาษาอังกฤษของคณะวิศวกรรมศาสตร์และคณะแพทยศาสตร์ให้ Gemini 2.5 Flash ประมวลผล พบว่า AI สามารถสร้างบทสรุปภาษาไทยและให้คำแนะนำขั้นตอนการตั้งค่าแก้ไข (เช่น การเพิ่ม Header ใน Nginx config) ได้อย่างถูกต้อง แม่นยำ และรวดเร็วภายในเวลาเฉลี่ยเพียง 2.1 วินาที')

    # Save
    doc.save(output_path)
    print("Chapter 4 docx created successfully at:", output_path)

if __name__ == '__main__':
    create_chapter4_document('/app/python/kku_scorecard_chapter4_final.docx')
